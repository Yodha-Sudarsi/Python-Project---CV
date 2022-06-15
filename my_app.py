
from docx import Document
from docx.shared import Inches
import pyttsx3

def speak (text):
    pyttsx3.speak(text)

document = Document()

# Profile photo
document.add_picture("hello.jpg",width = Inches(2.0))

# About name , phone number and email
name = input('What is your name ? ')
speak('Hello' + name + 'How are you today ?')
speak('What is your phone number  ? ')
phone_number = input('What is your phone number  ? ')
email = input('What is your email ? ')
document.add_paragraph(name + ' | ' + phone_number + ' | ' + email )  

# About me
document.add_heading('About me')
document.add_paragraph(input('Tell about yourself ? '))

# Work experience
document.add_heading('Work experience : ')
p = document.add_paragraph()
company = input('Enter company : ')
from_date = input('From date : ')
to_date = input('To date : ')
p.add_run(company + ' :  ').bold = True
p.add_run(from_date + '-' + to_date + '\n' ).italic = True

experience_details = input('describe your experiece at ' + company + ' ')
p.add_run(experience_details)

# More experiences
while True :
    has_more_experiences = input('Do you have more experiences ? Yes or No ')
    if has_more_experiences.lower() == 'yes' :
        p = document.add_paragraph()
        company = input('Enter company : ')
        from_date = input('From date : ')
        to_date = input('To date : ')
        p.add_run(company + ' :  ').bold = True
        p.add_run(from_date + '-' + to_date + '\n' ).italic = True

        experience_details = input('describe your experiece at ' + company + ' ')
        p.add_run(experience_details)
    else :
        break

# Skills
document.add_heading('Skills')
skill = input('Enter skill : ')
p=document.add_paragraph(skill)
p.style = 'List Bullet'

while True :
    has_more_skills = input('Do you have more skills ? Yes or No : ')
    if has_more_skills.lower() == 'yes' :
        skill = input('Enter skill : ')
        p=document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

# Footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using yashcode and Intuit quickbooks project"

document.save('cv.docx')
